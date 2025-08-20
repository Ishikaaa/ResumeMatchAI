# importing libraries
import os
import docx
from config.config_file import *
from utils.preprocessing import *
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import numpy as np
import faiss


def extract_docx_text_and_tables(file_path):
    """
    Extract text and tables from a .docx file, ignoring images.
    """
    doc = docx.Document(file_path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            cleaned_text = preprocess_text(para.text.strip())
            text_parts.append(cleaned_text)

    # Extract tables (only text, no images)
    for table in doc.tables:
        for row in table.rows:
            row_data = [
                preprocess_text(cell.text.strip())
                for cell in row.cells
            ]
            text_parts.append("\t".join(row_data))

    return "\n".join(text_parts)

def extract_text():
    # TODO - add doc
    resumes_data = {}

    for file_name in os.listdir(resume_folder_path):
        file_path = os.path.join(resume_folder_path, file_name)
        if file_name.lower().endswith(".docx"):
            text_data = extract_docx_text_and_tables(file_path)
            resumes_data[file_name] = text_data

        # TODO - work on .doc files

    return resumes_data


def chunking(resumes_data):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=300)
    texts = []
    metadatas = []

    # for name, content in resumes_data.items():
    #     chunks = splitter.split_text(content)
    #     texts.extend(chunks)
    #     metadatas.extend([{"source": name}] * len(chunks))

    for name, content in resumes_data.items():
        chunks = splitter.split_text(content)
        texts.extend(chunks)
        metadatas.extend([{"source": name}] * len(chunks))

    return texts, metadatas

# Step 3: Create embeddings
def get_embeddings(texts):
    """
    Create embeddings for given texts using HuggingFace embeddings model.
    """
    embedding_model = HuggingFaceEmbeddings(model_name=embedding_model_name)
    embeddings = embedding_model.embed_documents(texts)
    embeddings = np.array(embeddings).astype('float32')

    # Normalize embeddings → L2 norm = 1
    faiss.normalize_L2(embeddings)
    return embeddings, embedding_model

# Step 4: Store in FAISS
def store_in_faiss(texts, metadatas, embedding_model, index_path=faiss_path):
    """
    Store texts + metadata in FAISS using the provided embedding model.
    """
    # FAISS wrapper will take the embeddings already normalized
    print("💾 Storing in FAISS...")
    faiss_index = FAISS.from_texts(texts, embedding_model, metadatas=metadatas)
    faiss_index.save_local(index_path)
    print(f"Embeddings stored in FAISS at '{index_path}'")
    return faiss_index

def extract_skills_rake(jd_text, top_k=20):
    """
    Extract keywords/skills from job description text using RAKE.

    Args:
        jd_text (str): Job description text
        top_k (int): Number of top keywords to extract (optional)

    Returns:
        str: Keywords joined as a single string
    """
    r = Rake()
    r.extract_keywords_from_text(jd_text)
    # Get ranked phrases and clean unwanted characters
    phrases = [re.sub(r"[.)]", "", p) for p in r.get_ranked_phrases()[:top_k]]
    return " ".join(phrases)

if __name__ == "__main__":
    # Extract text nad preprocess them
    resumes_data = extract_text()

    # Step 2: Extract key skills from each resume before chunking
    for name, content in resumes_data.items():
        resume_skills = extract_skills_rake(content)
        resume_skills = remove_stopwords(resume_skills)
        resumes_data[name] = resume_skills

    # Chunking
    texts, metadatas = chunking(resumes_data)
    clean_texts = []
    for t in texts:
        t = deduplicate_words(t)
        clean_texts.append(t)

    # Step 3: Get embeddings model
    embeddings, embedding_model = get_embeddings(clean_texts)

    # Step 4: Store in FAISS
    store_in_faiss(texts, metadatas, embedding_model)
