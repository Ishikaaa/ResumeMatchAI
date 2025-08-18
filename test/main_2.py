# importing libraries
import os
import docx
import re
import json
from config.config_file import *
from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from transformers import AutoTokenizer

def extract_docx_text_and_tables(file_path):
    """
    Extract text and tables from a .docx file, ignoring images.
    """
    doc = docx.Document(file_path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # text_parts.append(para.text.strip())
            text_parts.append(preprocess_text(para.text.strip()))

    # Extract tables (only text, no images)
    for table in doc.tables:
        for row in table.rows:
            # row_data = [cell.text.strip() for cell in row.cells]
            row_data = [preprocess_text(cell.text.strip()) for cell in row.cells]
            text_parts.append("\t".join(row_data))

    return "\n".join(text_parts)


def extract_text():
    # TODO - add doc
    resumes_data = {}

    for file_name in os.listdir(resume_folder_path):
        file_path = os.path.join(resume_folder_path, file_name)
        if file_name.lower().endswith(".docx"):
            text_data = extract_docx_text_and_tables(file_path)
            resumes_data[file_name] = text_data

        # TODO - work on .doc files

    return resumes_data


def preprocess_text(text):
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)  # Remove extra spaces/newlines
    return text


tokenizer = AutoTokenizer.from_pretrained("google/flan-t5-base")

def chunk_text(text, max_tokens=512):
    tokens = tokenizer.encode(text)
    chunks = []
    for i in range(0, len(tokens), max_tokens):
        chunk = tokenizer.decode(tokens[i:i+max_tokens])
        chunks.append(chunk)
    return chunks


from transformers import pipeline

# Instruction-tuned model
extractor = pipeline("text2text-generation", model="google/flan-t5-large")


def extract_resume_info(resume_text):
    prompt = f"""
    Extract the following information from the resume below:
    - Skills
    - Certifications
    - Education
    - Work Experience (role, company, duration)
    Return in JSON format.

    Resume:
    {resume_text}
    """
    response = extractor(prompt, max_new_tokens=300, do_sample=False)[0]['generated_text']
    return response.strip()


def extract_full_resume_info(resume_text):
    chunks = chunk_text(resume_text, max_tokens=512)
    combined_info = {"Skills": [], "Certifications": [], "Education": [], "Experience": []}

    for chunk in chunks:
        info_json = extract_resume_info(chunk)
        try:
            info = json.loads(info_json)
            for key in combined_info:
                if key in info:
                    combined_info[key].extend(info[key] if isinstance(info[key], list) else [info[key]])
        except:
            # fallback if JSON parsing fails
            continue
    return combined_info

if __name__ == "__main__":
    # Step 1: Extract and preprocess resumes
    print("Extracting text from resumes...")
    resumes_data = extract_text()
    print(f"Total resumes found: {len(resumes_data)}\n")

    # Step 2: Extract structured info from each resume using LLM
    all_resumes_info = {}
    for resume_name, resume_text in resumes_data.items():
        print(f"Processing resume: {resume_name}")
        resume_info = extract_full_resume_info(resume_text)
        all_resumes_info[resume_name] = resume_info

        # Optional: print extracted info
        print(json.dumps(resume_info, indent=2))
        print("-" * 50)
