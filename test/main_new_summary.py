# # importing libraries
# import os
# import docx
# import re
# from config.config_file import *
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# # from langchain_community.embeddings import OpenAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from langchain_huggingface import HuggingFaceEmbeddings
# from langchain_community.vectorstores import FAISS
#
#
# def preprocess_text(text):
#     text = text.strip()  # remove leading/trailing spaces
#     text = text.lower()  # lowercase
#     text = re.sub(r'\s+', ' ', text)  # collapse multiple spaces/newlines
#     text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # remove non-ASCII characters
#     text = re.sub(r'[^\w\s]', '', text)  # remove punctuation
#     return text
#
# def extract_docx_text_and_tables(file_path):
#     """
#     Extract text and tables from a .docx file, ignoring images.
#     """
#     doc = docx.Document(file_path)
#     text_parts = []
#
#     # Extract paragraphs
#     for para in doc.paragraphs:
#         if para.text.strip():
#             # text_parts.append(para.text.strip())
#             # text_parts.append(preprocess_text(para.text.strip()))
#             cleaned_text = preprocess_text(para.text.strip())
#             # cleaned_text = deduplicate_words(cleaned_text)  # deduplicate words
#             text_parts.append(cleaned_text)
#
#     # Extract tables (only text, no images)
#     for table in doc.tables:
#         for row in table.rows:
#             # # row_data = [cell.text.strip() for cell in row.cells]
#             # row_data = [preprocess_text(cell.text.strip()) for cell in row.cells]
#             # text_parts.append("\t".join(row_data))
#             row_data = [
#                 preprocess_text(cell.text.strip())
#                 for cell in row.cells
#             ]
#             text_parts.append("\t".join(row_data))
#
#     return "\n".join(text_parts)
#
# def extract_text():
#     # TODO - add doc
#     resumes_data = {}
#
#     for file_name in os.listdir(resume_folder_path):
#         file_path = os.path.join(resume_folder_path, file_name)
#         if file_name.lower().endswith(".docx"):
#             text_data = extract_docx_text_and_tables(file_path)
#             resumes_data[file_name] = text_data
#
#         # TODO - work on .doc files
#     return resumes_data
#
# def chunking(resumes_data):
#     splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
#     texts = []
#     metadatas = []
#     summaries = []
#
#     # for name, content in resumes_data.items():
#     #     chunks = splitter.split_text(content)
#     #     texts.extend(chunks)
#     #     metadatas.extend([{"source": name}] * len(chunks))
#
#     for name, content in resumes_data.items():
#         chunks = splitter.split_text(content)
#         for chunk in chunks:
#             summary = resume_summary(chunk)  # summarize immediately
#             summaries.append(summary)
#             metadatas.append({"source": name})
#
#     return texts, metadatas
#
# def resume_summary(resume_text):
#     prompt = f"""
# check grammer as this I am giving to LLM model as prompt - You are an expert recruiter. Extract relevant skills, work experience, Education, Certifications, Projects details.
# Make sure to not miss even 1 skill from resume as this is very important for candiadate and for company
#
# Resume Text:
# {resume_text}
# """
#     result = summarizer(prompt, max_new_tokens=1024, temperature=0)
#     return result[0]["generated_text"]
#
# if __name__ == "__main__":
#     # Extract text nad preprocess them
#     resumes_data = extract_text()
#
#     texts, metadatas = chunking(resumes_data)
#
#     summarizer = pipeline(
#         "text2text-generation",
#         model="google/flan-t5-base",
#         tokenizer="google/flan-t5-base"
#     )
#




import os
from docx import Document
from transformers import pipeline
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS

# -------------------------
# Step 1: Load resumes and extract text
# -------------------------
def extract_text(folder_path="Resumes"):
    resumes = {}
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            doc = Document(os.path.join(folder_path, file_name))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            resumes[file_name] = full_text
    return resumes

# -------------------------
# Step 2: Initialize summarizer
# -------------------------
def get_summarizer():
    return pipeline(
        "text2text-generation",
        model="google/flan-t5-base",
        tokenizer="google/flan-t5-base"
    )

# -------------------------
# Step 3: Summarize a chunk
# -------------------------
def resume_summary(resume_text, summarizer):
    print("#### resume_text:", resume_text)
    prompt = f"""
    check grammer as this I am giving to LLM model as prompt - You are an expert recruiter. Extract relevant skills, work experience, Education, Certifications, Projects details. 
    Make sure to not miss even 1 skill from resume as this is very important for candiadate and for company

    Resume Text:
    {resume_text}
    """
    result = summarizer(prompt, max_new_tokens=1024, temperature=0)
    print(result[0]["generated_text"])
    return result[0]["generated_text"]
# -------------------------
# Step 4: Chunk text
# -------------------------
def chunk_text(text, chunk_size=500, chunk_overlap=50):
    # splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,  # 500 characters per chunk
        chunk_overlap=50,  # overlap to maintain context
        separators=["\n\n", "\n", ".", "!", "?"]  # try to split on paragraphs and sentences
    )
    return splitter.split_text(text)

# -------------------------
# Step 5: Embed and save FAISS index
# -------------------------
def create_faiss_index(texts, metadatas, index_path="resume_faiss_index_new"):
    embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    faiss_index = FAISS.from_texts(texts, embedding_model, metadatas=metadatas)
    faiss_index.save_local(index_path)
    print(f"FAISS index saved at '{index_path}'")

# -------------------------
# Step 6: Process all resumes
# -------------------------
def process_resumes(folder_path="Resumes", faiss_index_path="resume_faiss_index"):
    # Load resumes
    resumes_data = extract_text(folder_path)

    # Initialize summarizer
    summarizer = get_summarizer()

    all_texts = []
    all_metadatas = []

    # Process each resume
    for name, content in resumes_data.items():
        # print(f"@@@ name: {name}")
        # print(f"@@@ data type of content: {type(content)}")
        # print(f"@@@ length of content: {len(content)}")
        # print(f"@@@ content: {content}")
        # print(f"@@@ content: {content}")
        chunks = chunk_text(content)
        for chunk in chunks:
            print(f"@@@ data type of chunk: {type(chunk)}")
            print(f"@@@ length of chunk: {len(chunk)}")
            print(f"@@@ chunk: {chunk}")
            summary = resume_summary(chunk, summarizer)
            all_texts.append(summary)
            print(f"*********************************** summary: {summary}")

            all_metadatas.append({"source": name})
        break

    # Create FAISS index
    create_faiss_index(all_texts, all_metadatas, faiss_index_path)

# -------------------------
# Main
# -------------------------
if __name__ == "__main__":
    process_resumes(folder_path="../data/Resumes", faiss_index_path="../resume_faiss_index")
