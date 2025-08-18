# TODO - change file name

# importing libraries
import os
import docx
import re
from config.config_file import *
from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import numpy as np
import faiss


RESUME_STOPWORDS = [
    "implemented", "implementations", "developed", "worked", "working",
    "experience", "responsible", "handled", "performed", "participated", "involved"
]

def remove_resume_stopwords(text):
    words = text.split()
    filtered = [w for w in words if w not in RESUME_STOPWORDS]
    return " ".join(filtered)

def extract_docx_text_and_tables(file_path):
    """
    Extract text and tables from a .docx file, ignoring images.
    """
    doc = docx.Document(file_path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # text_parts.append(para.text.strip())
            # text_parts.append(preprocess_text(para.text.strip()))
            cleaned_text = preprocess_text(para.text.strip())
            cleaned_text = remove_resume_stopwords(cleaned_text)
            # cleaned_text = deduplicate_words(cleaned_text)  # deduplicate words
            text_parts.append(cleaned_text)

    # Extract tables (only text, no images)
    for table in doc.tables:
        for row in table.rows:
            # # row_data = [cell.text.strip() for cell in row.cells]
            # row_data = [preprocess_text(cell.text.strip()) for cell in row.cells]
            # text_parts.append("\t".join(row_data))
            row_data = [
                remove_resume_stopwords(preprocess_text(cell.text.strip()))
                for cell in row.cells
            ]
            text_parts.append("\t".join(row_data))

    return "\n".join(text_parts)

def extract_text():
    # TODO - add doc
    resumes_data = {}

    for file_name in os.listdir(resume_folder_path):
        file_path = os.path.join(resume_folder_path, file_name)
        if file_name.lower().endswith(".docx"):
            text_data = extract_docx_text_and_tables(file_path)
            resumes_data[file_name] = text_data

        # TODO - work on .doc files

    # # Example: print extracted content from files
    # for name, content in resumes_data.items():
    #     if '41' in name:
    #         print(name)
    #         print(content)
    #     # print(f"{name}:\n{content[:300]}\n{'-'*50}")

    return resumes_data

def deduplicate_words(text):
    """
    Remove repeated words, keep only unique words.
    """
    words = text.split()
    seen = set()
    unique_words = []
    for w in words:
        if w not in seen:
            seen.add(w)
            unique_words.append(w)
    return " ".join(unique_words)

def preprocess_text(text):
    # text = text.strip()
    # text = re.sub(r'\s+', ' ', text)  # Remove extra spaces/newlines
    # text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
    # return text.lower()  # Convert to lowercase
    #
    # text = text.lower()  # lowercase
    # text = re.sub(r'\s+', ' ', text)  # remove extra spaces
    # # text = re.sub(r'[^\w\s]', '', text)  # remove punctuation
    # return text.strip()

    text = text.strip()  # remove leading/trailing spaces
    text = text.lower()  # lowercase
    text = re.sub(r'\s+', ' ', text)  # collapse multiple spaces/newlines
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # remove non-ASCII characters
    text = re.sub(r'[^\w\s]', '', text)  # remove punctuation
    return text

def chunking(resumes_data):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    texts = []
    metadatas = []

    for name, content in resumes_data.items():
        chunks = splitter.split_text(content)
        texts.extend(chunks)
        metadatas.extend([{"source": name}] * len(chunks))

    return texts, metadatas

# Step 3: Create embeddings
def get_embeddings(texts):
    """
    Create embeddings for given texts using HuggingFace embeddings model.
    """
    # embedding_model = HuggingFaceEmbeddings(model_name=embedding_model_name)
    # embeddings = embedding_model.embed_documents(texts)
    # return embeddings, embedding_model

    embedding_model = HuggingFaceEmbeddings(model_name=embedding_model_name)
    embeddings = embedding_model.embed_documents(texts)
    embeddings = np.array(embeddings).astype('float32')

    # Normalize embeddings → L2 norm = 1
    faiss.normalize_L2(embeddings)
    return embeddings, embedding_model

# Step 4: Store in FAISS
def store_in_faiss(texts, metadatas, embedding_model, index_path=faiss_path):
    """
    Store texts + metadata in FAISS using the provided embedding model.
    """
    # print("💾 Storing in FAISS...")
    # faiss_index = FAISS.from_texts(texts, embedding_model, metadatas=metadatas)
    # faiss_index.save_local(index_path)  # persist index locally
    # print(f"Embeddings stored in FAISS at '{index_path}'")
    # return faiss_index

    # FAISS wrapper will take the embeddings already normalized
    print("💾 Storing in FAISS...")
    faiss_index = FAISS.from_texts(texts, embedding_model, metadatas=metadatas)
    faiss_index.save_local(index_path)
    print(f"Embeddings stored in FAISS at '{index_path}'")
    return faiss_index

if __name__ == "__main__":
    # Extract text nad preprocess them
    resumes_data = extract_text()
    # Chunking
    texts, metadatas = chunking(resumes_data)
    clean_texts = []
    for t in texts:
        # t = preprocess_text(t)
        # t = remove_resume_stopwords(t)
        t = deduplicate_words(t)
        clean_texts.append(t)

    # Step 3: Get embeddings model
    embeddings, embedding_model = get_embeddings(clean_texts)

    # Step 4: Store in FAISS
    store_in_faiss(texts, metadatas, embedding_model)

    # print(f"Total chunks created: {len(texts)}")
    # print(f"Example chunk:\n{texts[0]}")
    # print(metadatas)
    # print(len(metadatas))
