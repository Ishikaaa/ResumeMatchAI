import os
import docx

from config.config_file import *
from utils.preprocessing import *


def extract_docx_text_and_tables(file_path):
    """
    Extract text and tables from a .docx file, ignoring images.
    """
    doc = docx.Document(file_path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            cleaned_text = preprocess_text(para.text.strip())
            text_parts.append(cleaned_text)

    # Extract tables (only text, no images)
    for table in doc.tables:
        for row in table.rows:
            row_data = [
                preprocess_text(cell.text.strip())
                for cell in row.cells
            ]
            text_parts.append("\t".join(row_data))

    return "\n".join(text_parts)


def extract_text():
    """
    Extract text from all .docx resumes in the configured folder.
    """
    resumes_data = {}

    for file_name in os.listdir(RESUME_FOLDER):
        file_path = os.path.join(RESUME_FOLDER, file_name)
        if file_name.lower().endswith(".docx"):
            text_data = extract_docx_text_and_tables(file_path)
            resumes_data[file_name] = text_data

    return resumes_data
