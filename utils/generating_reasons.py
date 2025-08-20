import os

import docx
from dotenv import load_dotenv
from google import genai

from config.config_file import *
from utils.preprocessing import *

# Load variables from .env into environment
load_dotenv()

# Get the API key
gemini_api_key = os.getenv("GEMINI_API_KEY")


def extract_docx_text_and_tables(file_path):
    """
    Extract text and tables from a .docx file, ignoring images.
    """
    doc = docx.Document(file_path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            cleaned_text = preprocess_text(para.text.strip())
            text_parts.append(cleaned_text)

    # Extract tables (only text, no images)
    for table in doc.tables:
        for row in table.rows:
            row_data = [
                preprocess_text(cell.text.strip())
                for cell in row.cells
            ]
            text_parts.append("\t".join(row_data))

    return "\n".join(text_parts)


def extract_text(resume_path):
    if resume_path.lower().endswith(".docx"):
        resume_text_data = extract_docx_text_and_tables(resume_path)

    return resume_text_data


def generate_hiring_reason(jd_name, resume_name):
    reason = ""
    resume_path = os.path.join(RESUME_FOLDER, resume_name)
    jd_path = os.path.join(JOB_DESCRIPTION_FOLDER, jd_name)

    # extract resume data
    if not os.path.exists(resume_path):
        raise FileNotFoundError(f"Resume '{resume_name}' not found in {RESUME_FOLDER}")
    resumes_data = extract_text(resume_path)

    # extract JD data
    if not os.path.exists(jd_path):
        raise FileNotFoundError(f"Resume '{jd_name}' not found in {JOB_DESCRIPTION_FOLDER}")
    jd_data = extract_text(jd_path)

    # Ask model why candidate fits
    prompt = f"""
    Candidate Resume:
    {resumes_data}

    Job Description:
    {jd_data}

    Question: Why does this candidate fit for the job description? Give a clear and concise answer.
    Write a concise, bullet-point hiring recommendation (3–5 bullets) explaining why this candidate is a strong fit for the job description. Use recruiter-friendly language
    """

    client = genai.Client()

    response = client.models.generate_content(
        model="gemini-2.5-flash", contents=prompt
    )

    # Safely extract text or return empty string
    if hasattr(response, "text") and response.text:
        reason = response.text
    elif hasattr(response, "result") and response.result.candidates:
        reason = response.result.candidates[0].content.get("text", "")
    else:
        reason = "Candidate is a good fit for the role based on their resume and job description."

    return reason
